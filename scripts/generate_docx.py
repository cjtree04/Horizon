#!/usr/bin/env python3
"""
Convert Horizon Chinese daily summary to formatted Word document on Desktop.

Usage:
    uv run python scripts/generate_docx.py
    uv run python scripts/generate_docx.py --x-posts-file /tmp/x_posts.md
"""

import re
import sys
import argparse
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

HORIZON_DIR = Path(__file__).parent.parent
SUMMARIES_DIR = HORIZON_DIR / "data" / "summaries"
DESKTOP = Path.home() / "Desktop"


def find_latest_zh() -> Path:
    files = sorted(SUMMARIES_DIR.glob("horizon-*-zh.md"), reverse=True)
    if not files:
        raise FileNotFoundError(f"No zh summary found in {SUMMARIES_DIR}")
    return files[0]


def clean_md_links(text: str) -> str:
    return re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', text)


def clean_html(text: str) -> str:
    return re.sub(r'<[^>]+>', '', text)


def parse_items(content: str) -> list:
    items = []
    blocks = re.split(r'\n---\n', content)
    for block in blocks:
        if '<a id="item-' not in block:
            continue
        item = {}

        # Title + URL + score
        m = re.search(r'## \[(.+?)\]\((.+?)\)\s*⭐️\s*([\d.]+)/10', block)
        if not m:
            m = re.search(r'## (.+?)\s*⭐️\s*([\d.]+)/10', block)
            if not m:
                continue
            item['title'] = m.group(1).strip()
            item['url'] = ''
            item['score'] = float(m.group(2))
        else:
            item['title'] = m.group(1).strip()
            item['url'] = m.group(2).strip()
            item['score'] = float(m.group(3))

        # Source line (has · separator and known source keyword)
        src_m = re.search(
            r'\n((?:hackernews|reddit|rss|twitter|github)[^\n]+)', block, re.IGNORECASE
        )
        if src_m:
            src = clean_md_links(src_m.group(1)).strip()
            item['source'] = src
        else:
            item['source'] = ''

        # Main summary: text after source line, before first **label**
        after = block[src_m.end():] if src_m else block
        body_m = re.match(
            r'\s*\n+(.*?)(?=\n\n\*\*|\n\n<details|\Z)', after, re.DOTALL
        )
        if body_m:
            body = clean_html(body_m.group(1).strip())
            item['summary'] = body if not body.startswith('**') else ''
        else:
            item['summary'] = ''

        # 背景
        bg_m = re.search(r'\*\*背景\*\*[：:]\s*(.+?)(?=\n\n\*\*|\n\n<details|\Z)', block, re.DOTALL)
        item['background'] = clean_html(bg_m.group(1).strip()) if bg_m else ''

        # 社区讨论
        disc_m = re.search(r'\*\*社区讨论\*\*[：:]\s*(.+?)(?=\n\n\*\*标签|\Z)', block, re.DOTALL)
        item['discussion'] = clean_html(disc_m.group(1).strip()) if disc_m else ''

        # 标签
        tags_m = re.search(r'\*\*标签\*\*[：:]?\s*(.+)', block)
        if tags_m:
            raw = tags_m.group(1).strip()
            item['tags'] = re.sub(r'`#?([^`]+)`', r'#\1', raw)
        else:
            item['tags'] = ''

        items.append(item)
    return items


def score_color(score: float) -> RGBColor:
    if score >= 8.0:
        return RGBColor(27, 94, 32)    # green
    elif score >= 7.0:
        return RGBColor(230, 81, 0)    # orange
    return RGBColor(66, 66, 66)        # gray


def add_label_content(doc, label: str, content: str, label_color: RGBColor):
    """Add an indented label + content paragraph pair."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.6)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    lr = p.add_run(label)
    lr.bold = True
    lr.font.size = Pt(10)
    lr.font.color.rgb = label_color

    p2 = doc.add_paragraph()
    p2.paragraph_format.left_indent = Cm(0.8)
    p2.paragraph_format.space_after = Pt(4)
    r = p2.add_run(content)
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(60, 60, 60)


def add_divider(doc):
    p = doc.add_paragraph('─' * 52)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    for run in p.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(200, 200, 200)


def build_docx(date: str, meta: str, items: list, x_posts_content: str = None) -> Path:
    doc = Document()

    for section in doc.sections:
        section.left_margin = Cm(2.8)
        section.right_margin = Cm(2.8)
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)

    # ── 标题 ──
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run('AI 科技日报')
    r.bold = True
    r.font.size = Pt(26)
    r.font.color.rgb = RGBColor(21, 67, 96)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(date)
    r.font.size = Pt(13)
    r.font.color.rgb = RGBColor(90, 90, 90)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(14)
    r = p.add_run(meta)
    r.font.size = Pt(10)
    r.italic = True
    r.font.color.rgb = RGBColor(130, 130, 130)

    add_divider(doc)

    # ── 目录 ──
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run('今日要点')
    r.bold = True
    r.font.size = Pt(14)
    r.font.color.rgb = RGBColor(21, 67, 96)

    for i, item in enumerate(items, 1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.4)
        p.paragraph_format.space_after = Pt(3)
        num_r = p.add_run(f'{i}.  ')
        num_r.font.size = Pt(11)
        num_r.font.color.rgb = RGBColor(100, 100, 100)
        t_r = p.add_run(item['title'])
        t_r.font.size = Pt(11)
        p.add_run('  ')
        sc_r = p.add_run(f"⭐ {item['score']}")
        sc_r.font.size = Pt(10)
        sc_r.font.color.rgb = score_color(item['score'])

    if x_posts_content:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.4)
        p.paragraph_format.space_after = Pt(3)
        num_r = p.add_run(f'{len(items) + 1}.  ')
        num_r.font.size = Pt(11)
        num_r.font.color.rgb = RGBColor(100, 100, 100)
        t_r = p.add_run('X 热帖精选（小红书候选）')
        t_r.font.size = Pt(11)

    doc.add_page_break()

    # ── 正文条目 ──
    for i, item in enumerate(items, 1):
        # 标题行
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8 if i > 1 else 0)
        p.paragraph_format.space_after = Pt(4)
        nr = p.add_run(f'{i}. ')
        nr.bold = True
        nr.font.size = Pt(15)
        nr.font.color.rgb = RGBColor(21, 67, 96)
        tr = p.add_run(item['title'])
        tr.bold = True
        tr.font.size = Pt(14)

        # 评分 + 来源
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(5)
        sc_r = p.add_run(f"⭐ {item['score']}/10  ")
        sc_r.font.size = Pt(11)
        sc_r.bold = True
        sc_r.font.color.rgb = score_color(item['score'])
        if item['source']:
            src_r = p.add_run(item['source'])
            src_r.font.size = Pt(9)
            src_r.font.color.rgb = RGBColor(140, 140, 140)

        # 摘要
        if item['summary']:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.4)
            p.paragraph_format.space_after = Pt(6)
            r = p.add_run(item['summary'])
            r.font.size = Pt(11)

        # 背景
        if item['background']:
            add_label_content(doc, '背景知识：', item['background'], RGBColor(31, 97, 141))

        # 社区讨论
        if item['discussion']:
            add_label_content(doc, '社区讨论：', item['discussion'], RGBColor(30, 130, 76))

        # 标签
        if item['tags']:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.4)
            p.paragraph_format.space_after = Pt(4)
            r = p.add_run(item['tags'])
            r.font.size = Pt(9)
            r.italic = True
            r.font.color.rgb = RGBColor(160, 160, 160)

        # 链接
        if item['url']:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.4)
            p.paragraph_format.space_after = Pt(10)
            r = p.add_run(f"🔗 {item['url']}")
            r.font.size = Pt(9)
            r.font.color.rgb = RGBColor(41, 128, 185)

        add_divider(doc)

    # ── X 热帖精选 ──
    if x_posts_content:
        doc.add_page_break()

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(10)
        r = p.add_run('X 热帖精选')
        r.bold = True
        r.font.size = Pt(20)
        r.font.color.rgb = RGBColor(21, 67, 96)

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(16)
        r = p.add_run('小红书候选 · 供筛选')
        r.font.size = Pt(11)
        r.italic = True
        r.font.color.rgb = RGBColor(130, 130, 130)

        add_divider(doc)
        doc.add_paragraph()

        # 渲染 x_posts_content（纯文本，按段落输出）
        for line in x_posts_content.strip().split('\n'):
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
            r = p.add_run(line)
            # 给编号行加粗
            if re.match(r'^\d+[\.\、]', line.strip()):
                r.bold = True
                r.font.size = Pt(12)
                r.font.color.rgb = RGBColor(21, 67, 96)
            else:
                r.font.size = Pt(11)

    out = DESKTOP / f"AI日报-{date}.docx"
    doc.save(str(out))
    return out


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument('--x-posts-file', default=None,
                        help='Path to X posts markdown file (generated by Track 2)')
    args = parser.parse_args()

    x_posts = None
    if args.x_posts_file:
        x_path = Path(args.x_posts_file)
        if x_path.exists():
            x_posts = x_path.read_text(encoding='utf-8')

    zh_md = find_latest_zh()
    content = zh_md.read_text(encoding='utf-8')

    date_m = re.search(r'(\d{4}-\d{2}-\d{2})', content)
    date = date_m.group(1) if date_m else datetime.now().strftime('%Y-%m-%d')

    meta_m = re.search(r'> (.+)', content)
    meta = clean_md_links(meta_m.group(1)) if meta_m else ''

    items = parse_items(content)
    if not items:
        print('❌ 未找到有效条目，请检查 zh 日报格式')
        sys.exit(1)

    out = build_docx(date, meta, items, x_posts)
    print(f'✅ 已生成：{out}')


if __name__ == '__main__':
    main()
